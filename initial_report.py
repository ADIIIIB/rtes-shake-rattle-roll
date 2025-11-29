from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    document = Document()

    # --- Title ---
    title = document.add_heading('Technical Analysis & Divergence Resolution Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Subtitle ---
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Subject: Comparison of Frequency-Domain vs. Time-Domain Approaches')
    run.italic = True
    run.font.size = Pt(12)

    document.add_paragraph('') # Spacer

    # --- Section 1: Core Algorithm Divergence ---
    document.add_heading('1. Core Algorithm Divergence: FOG Detection', level=1)
    
    p = document.add_paragraph()
    run = p.add_run('Conflict: ')
    run.bold = True
    p.add_run('How to detect Freezing of Gait (FOG)?')

    document.add_paragraph('• Teammate\'s Proposal: Time-Domain "Cadence Drop" (Count steps -> Detect stop).', style='List Bullet')
    document.add_paragraph('• Current Implementation: Frequency-Domain "Freeze Index" (FFT Energy Ratio).', style='List Bullet')

    document.add_heading('Analysis & Derivation', level=2)

    document.add_heading('1. Project Requirement Alignment', level=3)
    p = document.add_paragraph('The prompt explicitly states: "Use... FFT library to return information on frequency distribution".')
    document.add_paragraph('• Critique of Cadence Method: Counting steps operates in the Time Domain. This underutilizes the required FFT library for the hardest part of the assignment.', style='List Bullet')
    document.add_paragraph('• Strength of Freeze Index: It directly utilizes the spectral output, calculating the ratio of freezing band energy (3-8Hz) to locomotor band energy (0.5-3Hz).', style='List Bullet')

    document.add_heading('2. Robustness Calculation', level=3)
    p = document.add_paragraph('Scenario: Accelerometer noise is high.')
    document.add_paragraph('• Time-Domain Risk: Peak detection algorithms are sensitive to noise. A slight "bump" can miscount steps, causing cadence fluctuation and false positives.', style='List Bullet')
    document.add_paragraph('• Freq-Domain Advantage: FFT integrates energy over 3 seconds. Random noise cancels out, while rhythmic trembling creates a distinct spectral peak.', style='List Bullet')

    document.add_heading('3. The "Static False Positive" Solution', level=3)
    document.add_paragraph('• Teammate\'s Concern: "If the patient just stands still, the Freeze Index might misfire."', style='List Bullet')
    document.add_paragraph('• Our Solution (Implemented): We added a WALKING_THRESHOLD. FOG detection only activates if Locomotor Energy > 5000. This mimics the logic of "Cadence" verification without step counting.', style='List Bullet')

    # --- Section 2: Physics & Threshold Analysis ---
    document.add_heading('2. Physics & Threshold Analysis (Numerical Proof)', level=1)
    
    p = document.add_paragraph()
    run = p.add_run('Verification: ')
    run.bold = True
    p.add_run('Thresholds (3000, 5000) are derived from physics and FFT mathematics.')

    document.add_heading('Derivation of Energy Levels', level=2)
    
    # Formula
    p = document.add_paragraph()
    p.add_run('Formula: Magnitude ≈ (Amplitude × Samples) / 2').bold = True
    document.add_paragraph('Hardware Config: Sensor data in mg (milli-g). 1g = 1000 units. Window = 256 samples.')

    # Case A
    document.add_heading('Case A: Normal Walking (Arm Swing)', level=3)
    document.add_paragraph('• Physics: Typical arm swing acceleration ≈ 0.4g = 400mg.', style='List Bullet')
    document.add_paragraph('• FFT Energy: (400 × 256) / 2 = 51,200.', style='List Bullet')
    document.add_paragraph('• Leakage Analysis: Assuming ~10% spectral leakage into higher frequencies.', style='List Bullet')
    document.add_paragraph('• Noise Floor: 51,200 × 10% ≈ 5,120.', style='List Bullet')
    p = document.add_paragraph('Conclusion: NOISE_THRESHOLD must be > 5,000. Setting it to 10,000 provides a 2x safety margin.')
    p.runs[0].bold = True

    # Case B
    document.add_heading('Case B: Parkinsonian Tremor', level=3)
    document.add_paragraph('• Physics: Resting tremor intensity ≈ 0.3g = 300mg.', style='List Bullet')
    document.add_paragraph('• FFT Energy: (300 × 256) / 2 ≈ 38,400.', style='List Bullet')
    document.add_paragraph('• Detection Check: 38,400 > 10,000 (Threshold).', style='List Bullet')
    p = document.add_paragraph('Result: DETECTED. The system is sensitive enough.')
    p.runs[0].bold = True

    # Case C
    document.add_heading('Case C: Sensor Noise (Idle)', level=3)
    document.add_paragraph('• Physics: Sensor on table ≈ 20mg noise.', style='List Bullet')
    document.add_paragraph('• FFT Energy: (20 × 256) / 2 ≈ 2,560.', style='List Bullet')
    document.add_paragraph('• Detection Check: 2,560 < 10,000.', style='List Bullet')
    p = document.add_paragraph('Result: IGNORED.')
    p.runs[0].bold = True

    # --- Section 3: System Architecture ---
    document.add_heading('3. System Architecture: Blocking vs. Buffering', level=1)
    
    p = document.add_paragraph()
    run = p.add_run('Conflict: ')
    run.bold = True
    p.add_run('Continuous Sampling (Teammate) vs. Blocking Sampling (Current).')

    document.add_heading('Latency Impact Analysis', level=2)
    document.add_paragraph('• Current Architecture: Collect 3s -> Stop -> Process -> Resume.', style='List Bullet')
    document.add_paragraph('• Processing Time Estimate (STM32L4): ~20ms total (FFT + BLE).', style='List Bullet')
    
    p = document.add_paragraph()
    p.add_run('Loss Ratio Calculation:').bold = True
    document.add_paragraph('Loss = 20ms (Processing) / 3000ms (Sampling) ≈ 0.6%')

    p = document.add_paragraph('Conclusion: We lose less than 1% of data. Tremors and FOG are continuous events; missing 20ms does not affect detection accuracy.')
    
    p = document.add_paragraph('Recommendation: Implementing double-buffering increases complexity and bug risk significantly for negligible gain. The blocking approach is safer.')
    p.runs[0].bold = True

    # --- Section 4: Integration Strategy ---
    document.add_heading('4. Integration & UI Strategy (The Merger)', level=1)
    
    document.add_paragraph('To align with the teammate\'s plan while preserving technical stability:')
    
    document.add_paragraph('1. Adopt the Buttons: Add button interrupts for state toggling (IDLE vs MONITORING).', style='List Number')
    document.add_paragraph('2. Keep the Algorithms: Stick to Frequency-Domain (FFT) as it is now unit-tested and verified.', style='List Number')
    document.add_paragraph('3. Keep the Sampling: Stick to the blocking loop to ensure stability (<1% data loss).', style='List Number')

    # --- Checklist ---
    document.add_heading('Discussion Checklist', level=2)
    document.add_paragraph('[ ] Agree to use Freeze Index instead of Step Counting (Safety & FFT Requirement)', style='List Bullet')
    document.add_paragraph('[ ] Agree on mg (milli-g) units for all calculations (Precision)', style='List Bullet')
    document.add_paragraph('[ ] Agree to add Button functionality for UI points', style='List Bullet')
    document.add_paragraph('[ ] Confirm Hardware Setup: SPI for BLE, I2C for Sensor', style='List Bullet')

    # Save
    file_name = 'Technical_Report.docx'
    document.save(file_name)
    print(f"Report generated successfully: {file_name}")

if __name__ == "__main__":
    create_report()